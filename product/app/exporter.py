from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from .config import EXPORTS_DIR


FORMATS = {"md", "docx", "pdf"}


def export_report(report: dict[str, Any], file_format: str) -> Path:
    markdown = _report_markdown(report)
    return _export(markdown, report["title"], f"report-{report['id']}", file_format)


def export_conversation(conversation: dict[str, Any], file_format: str) -> Path:
    lines = [
        f"# {conversation['title']}",
        "",
        f"- 对话来源：{conversation['source']}",
        f"- 创建时间：{conversation['created_at']}",
        f"- 最后更新：{conversation['updated_at']}",
        "",
    ]
    if conversation.get("company_name"):
        lines.insert(
            3,
            f"- 关联公司：{conversation['company_name']}（{conversation.get('company_market')} · "
            f"{conversation.get('company_symbol')}）",
        )
    for message in conversation.get("messages", []):
        role = "老板" if message["role"] == "user" else "AI 投研员工"
        lines.extend(
            [
                f"## {role} · {message['created_at']}",
                "",
                message["content"].strip(),
                "",
            ]
        )
        if message.get("sources"):
            lines.append("### 来源")
            lines.append("")
            lines.extend(
                f"- [{source.get('title') or '原始来源'}]({source.get('url')})"
                for source in message["sources"]
                if source.get("url")
            )
            lines.append("")
    return _export(
        "\n".join(lines),
        conversation["title"],
        f"conversation-{conversation['id']}",
        file_format,
    )


def _report_markdown(report: dict[str, Any]) -> str:
    source_lines = "\n".join(
        f"- **{item.get('quality_label') or '待核验'}｜"
        f"{item.get('citation_role') or '检索参考'}** "
        f"[{item.get('title') or '原始来源'}]({item.get('url')})"
        f"（{item.get('publisher') or item.get('domain') or '未知机构'}）"
        for item in report.get("sources", [])
        if item.get("url")
    )
    audit = report.get("source_audit") or {}
    warning_lines = "\n".join(f"- {item}" for item in audit.get("warnings", []))
    return (
        f"# {report['title']}\n\n"
        f"- 生成时间：{report['created_at']}\n"
        f"- 类型：{report['report_type']}\n"
        f"- 研究引擎：{report.get('engine') or '未记录'}\n"
        f"- 复核方式：{'事实核验员 + 反方研究员' if report.get('review_mode') == 'team' else '主研究员单独完成'}\n"
        f"- 模型：{report.get('model') or '未记录'}\n\n"
        f"{report['content'].strip()}\n\n"
        f"## 证据质量\n\n"
        f"- 结论：{audit.get('coverage_label', '尚未审计')}\n"
        f"- 一手来源：{audit.get('primary_count', 0)} / {audit.get('total', 0)}\n"
        f"- 正文引用：{audit.get('cited_count', 0)}\n"
        f"- 独立域名：{audit.get('unique_domains', 0)}\n"
        f"- 数字事实同行引用：{audit.get('cited_numeric_claim_count', 0)} / "
        f"{audit.get('numeric_claim_count', 0)}\n"
        f"{warning_lines}\n\n"
        f"## 来源清单\n\n{source_lines or '- 本报告没有联网来源。'}\n\n"
        "研究辅助，不构成投资建议。\n"
    )


def _export(markdown: str, title: str, key: str, file_format: str) -> Path:
    if file_format not in FORMATS:
        raise ValueError("仅支持 md、docx 或 pdf。")
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", title).strip("-")[:55] or key
    path = EXPORTS_DIR / f"{key}_{safe_title}.{file_format}"
    if file_format == "md":
        path.write_text(markdown, encoding="utf-8")
    elif file_format == "docx":
        _write_docx(path, markdown, title)
    else:
        _write_pdf(path, markdown, title)
    return path


def _write_docx(path: Path, markdown: str, title: str) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = title
    for line in markdown.splitlines():
        clean = _plain(line)
        if not clean:
            document.add_paragraph("")
        elif line.startswith("# "):
            document.add_heading(clean, level=1)
        elif line.startswith("## "):
            document.add_heading(clean, level=2)
        elif line.startswith("### "):
            document.add_heading(clean, level=3)
        elif re.match(r"^\s*[-*]\s+", line):
            document.add_paragraph(clean, style="List Bullet")
        elif re.match(r"^\s*\d+[.)]\s+", line):
            document.add_paragraph(clean, style="List Number")
        else:
            document.add_paragraph(clean)
    for style in document.styles:
        if hasattr(style, "font"):
            style.font.name = "PingFang SC"
            style.font.size = Pt(10.5)
    document.save(path)


def _write_pdf(path: Path, markdown: str, title: str) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9.5,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    heading = ParagraphStyle(
        "ChineseHeading",
        parent=body,
        fontSize=16,
        leading=22,
        spaceAfter=10,
    )
    subheading = ParagraphStyle(
        "ChineseSubheading",
        parent=body,
        fontSize=12,
        leading=18,
        spaceBefore=8,
        spaceAfter=5,
    )
    story = []
    for line in markdown.splitlines():
        clean = _plain(line)
        if not clean:
            story.append(Spacer(1, 3 * mm))
            continue
        style = heading if line.startswith("# ") else subheading if line.startswith("##") else body
        story.append(Paragraph(_xml_escape(clean), style))
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=17 * mm,
        bottomMargin=17 * mm,
        title=title,
    )
    document.build(story)


def _plain(line: str) -> str:
    value = re.sub(r"^\s*#{1,6}\s+", "", line)
    value = re.sub(r"^\s*[-*]\s+", "• ", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", value)
    value = value.replace("**", "").replace("__", "").replace("`", "")
    return value.strip()


def _xml_escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
